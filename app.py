# app.py — ADI Builder (robust assets path, clean UI, logo, courses, verbs, MCQs)
# Place `assets/` next to this file OR one/two levels above. You can also set ASSETS_DIR env var.
#   assets/
#     ├─ adi-logo.png     (optional)
#     └─ courses.csv      (code,label)  OR courses.json ([{"code":"..","label":".."}])

import os
import base64
import csv
import json
from io import StringIO, BytesIO
from pathlib import Path

import streamlit as st

# =========================
# Robust assets directory
# =========================
BASE_DIR = Path(__file__).resolve().parent

def resolve_assets_dir() -> Path:
    # 1) Env var wins if valid
    env = os.getenv("ASSETS_DIR")
    if env:
        p = Path(env).expanduser().resolve()
        if p.exists():
            return p

    # 2) Try beside app.py, then up one, then up two
    candidates = [
        BASE_DIR / "assets",
        BASE_DIR.parent / "assets",
        BASE_DIR.parent.parent / "assets",
        Path.cwd() / "assets",
    ]
    for c in candidates:
        if c.exists():
            return c.resolve()

    # 3) Fallback to beside app.py (even if missing) to keep paths consistent
    return (BASE_DIR / "assets").resolve()

ASSETS_DIR = resolve_assets_dir()

# =========================
# Page & Theme
# =========================
st.set_page_config(page_title="ADI Builder — Lesson Activities & Questions", layout="wide")

ADI_GREEN = "#245a34"
STONE     = "#F5F4F2"
MUTED     = "#6b7280"

# Set True to hide the "Browse files" buttons in uploaders (drag&drop only)
HIDE_BROWSE_BUTTON = False

st.markdown(f"""
<style>
  .block-container {{ max-width: 980px; margin: 0 auto; padding-top: .6rem; }}

  /* Sticky premium banner */
  .adi-topbar {{
    position: sticky; top: 0; z-index: 1000;
    display:flex; align-items:center; gap:12px;
    padding:.55rem .9rem; background:{ADI_GREEN}; color:#fff;
    border-radius:0 0 12px 12px;
    box-shadow: 0 1px 0 rgba(0,0,0,.05);
  }}
  .adi-topbar img {{ height:32px; display:block; }}
  .adi-title {{ font-weight:700; letter-spacing:.2px; margin:0; font-size:1.06rem; }}
  .adi-badge {{
    display:inline-flex; align-items:center; justify-content:center;
    width:28px; height:28px; border-radius:50%;
    background:rgba(255,255,255,.18); font-weight:700; font-size:.9rem;
  }}

  .adi-card {{ background:{STONE}; border:1px solid #e7e5e4; border-radius:12px; padding:12px; }}
  .muted {{ color:{MUTED}; }}

  .tight > div {{ margin-bottom:.35rem; }}
  .stButton > button {{ border-radius:10px; }}

  /* Segmented control: active = ADI green */
  [data-baseweb="segmented-control"] [aria-selected="true"] {{
    background:{ADI_GREEN} !important; color:white !important;
  }}

  /* Strong labels */
  label:has(+ div [role="listbox"]),
  label:has(+ div input[type="number"]),
  label:has(+ div input[type="text"]) {{
    font-weight: 600 !important; color: #374151 !important;
  }}

  /* Pills */
  .pill {{ display:inline-block; padding:2px 10px; border-radius:999px; font-size:.85rem;
          line-height:1.6; white-space:nowrap; margin-left:6px; }}
  .pill-green {{ background:{ADI_GREEN}; color:#fff; }}

  /* Visible dashed dropzone */
  .dropzone-visible {{
    border:2.5px dashed {ADI_GREEN}; background:#f6faf7;
    border-radius:12px; padding:.6rem .6rem;
  }}
  [data-testid="stFileUploaderDropzone"],
  div[aria-label="File dropzone"],
  div:has(> input[type="file"]) {{
    border:2.5px dashed {ADI_GREEN} !important;
    background:#f6faf7 !important;
    border-radius:12px !important;
    transition: border-color .15s ease, background .15s ease;
  }}
  [data-testid="stFileUploaderDropzone"]:hover,
  div[aria-label="File dropzone"]:hover {{
    border-color:#1e4b2b !important;
    background:#eef7f1 !important;
  }}
  [data-testid="stFileUploaderDropzone"] p,
  div[aria-label="File dropzone"] p {{
    color:#0f3d22 !important; margin-bottom:0 !important;
  }}

  /* Upload chips / notes */
  .upload-chip {{
    display:inline-flex; align-items:center; gap:.4rem;
    padding:2px 10px; border-radius:999px;
    background:{ADI_GREEN}; color:#fff; font-size:.85rem; font-weight:600;
  }}
  .upload-chip svg {{ width:14px; height:14px; display:block; }}
  .upload-note {{
    border-left:4px solid {ADI_GREEN}; background:#f4fbf6;
    padding:.5rem .75rem; border-radius:8px; color:#0f3d22; margin-top:.5rem;
  }}
  .upload-warn {{
    border-left:4px solid #b45309; background:#fff7ed;
    padding:.5rem .75rem; border-radius:8px; color:#7c2d12; margin-top:.5rem;
  }}
  .soft-tip {{
    border-left:4px solid #eab308; background:#fffbeb;
    padding:.5rem .75rem; border-radius:8px; color:#713f12; margin:.6rem 0 0 0;
  }}
</style>
""", unsafe_allow_html=True)

if HIDE_BROWSE_BUTTON:
    st.markdown("""
    <style>
      [data-testid="stFileUploader"] button { display: none !important; }
    </style>
    """, unsafe_allow_html=True)

# =========================
# Helpers
# =========================
def b64_bytes(b: bytes) -> str:
    return base64.b64encode(b).decode("utf-8")

def b64_file(path: Path) -> str | None:
    try:
        return b64_bytes(path.read_bytes())
    except Exception:
        return None

def make_courses_template() -> bytes:
    rows = [
        ("GE4-EPM","Defense Technology Practices"),
        ("GE4-IPM","Integrated Project & Materials Mgmt"),
        ("GE4-MRO","Military Vehicle & Aircraft MRO"),
        ("CT4-COM","Computation for Chemical Technologists"),
        ("CT4-EMG","Explosives Manufacturing"),
        ("CT4-TFL","Thermofluids"),
    ]
    s = StringIO(); w = csv.writer(s)
    w.writerow(["code","label"])
    for r in rows: w.writerow(r)
    return s.getvalue().encode("utf-8")

# =========================
# Session
# =========================
def init_state():
    ss = st.session_state
    ss.setdefault("course_code", "")
    ss.setdefault("class_cohort", "D1-C01")
    ss.setdefault("lesson", 1)
    ss.setdefault("week", 1)
    ss.setdefault("instructor", "Daniel")
    ss.setdefault("topic_outcome", "")
    ss.setdefault("mode", "Knowledge")
    ss.setdefault("topics_text", "Topic A\nTopic B\nTopic C")
    ss.setdefault("bloom_level", "Low")
    ss.setdefault("verbs_selected", [])
    ss.setdefault("generated_items", [])
    ss.setdefault("COURSES", None)
    ss.setdefault("logo_b64", None)
    ss.setdefault("logo_uploaded", False)
    ss.setdefault("logo_file_info", {})
    ss.setdefault("logo_warning", "")
    ss.setdefault("courses_uploaded", False)
    ss.setdefault("courses_file_info", {})
init_state()

# =========================
# Courses (assets override or fallback)
# =========================
def load_courses_from_assets() -> list[tuple[str,str]]:
    items: list[tuple[str,str]] = []
    csvp = ASSETS_DIR / "courses.csv"
    jsp  = ASSETS_DIR / "courses.json"

    if csvp.exists():
        with csvp.open("r", encoding="utf-8") as f:
            for r in csv.DictReader(f):
                code = (r.get("code") or "").strip()
                label = (r.get("label") or "").strip()
                if code and label:
                    items.append((code, label))
    elif jsp.exists():
        raw = json.loads(jsp.read_text(encoding="utf-8"))
        for r in raw:
            code = (r.get("code") or "").strip()
            label = (r.get("label") or "").strip()
            if code and label:
                items.append((code, label))

    if items:
        return items

    # Fallback small list
    return [
        ("GE4-EPM","Defense Technology Practices"),
        ("GE4-IPM","Integrated Project & Materials Mgmt"),
        ("GE4-MRO","Military Vehicle & Aircraft MRO"),
        ("CT4-COM","Computation for Chemical Technologists"),
        ("CT4-EMG","Explosives Manufacturing"),
        ("CT4-TFL","Thermofluids"),
    ]

if st.session_state.COURSES is None:
    st.session_state.COURSES = load_courses_from_assets()

HAS_ASSET_COURSES = (ASSETS_DIR / "courses.csv").exists() or (ASSETS_DIR / "courses.json").exists()

def set_courses(new_list: list[tuple[str,str]]):
    st.session_state.COURSES = new_list

def course_codes() -> list[str]:
    return [c for c,_ in st.session_state.COURSES]

def code_to_label() -> dict:
    return dict(st.session_state.COURSES)

# =========================
# Logo & Banner
# =========================
def resolve_logo_b64() -> str | None:
    if st.session_state.logo_b64:
        return st.session_state.logo_b64
    return b64_file(ASSETS_DIR / "adi-logo.png")

def render_topbar(logo_b64: str | None):
    logo_html = f'<img src="data:image/png;base64,{logo_b64}"/>' if logo_b64 else '<div class="adi-badge">A</div>'
    st.markdown(
        f'<div class="adi-topbar">{logo_html}<h1 class="adi-title">ADI Builder — Lesson Activities & Questions</h1></div>',
        unsafe_allow_html=True
    )

render_topbar(resolve_logo_b64())

# =========================
# Inline logo prompt (only if no logo yet)
# =========================
no_logo_anywhere = resolve_logo_b64() is None
if no_logo_anywhere:
    st.markdown("#### Add your logo")
    st.caption("Drop a PNG, JPG, or SVG to brand the top banner.")
    st.markdown("<div class='dropzone-visible'>", unsafe_allow_html=True)
    logo_inline = st.file_uploader("Drag & drop logo here", key="logo_inline")
    st.markdown("</div>", unsafe_allow_html=True)

    if logo_inline is not None:
        ext = Path(logo_inline.name).suffix.lower().lstrip(".")
        allowed_ext = {"png", "jpg", "jpeg", "svg"}
        if ext in allowed_ext:
            st.session_state.logo_b64 = base64.b64encode(logo_inline.getvalue()).decode("utf-8")
            st.session_state.logo_file_info = {"name": logo_inline.name, "size": logo_inline.size}
            st.session_state.logo_uploaded = True
            render_topbar(st.session_state.logo_b64)
            no_logo_anywhere = False
        else:
            st.markdown(
                f"<div class='upload-warn'>⚠️ Only PNG, JPG or SVG are supported for the banner logo. "
                f"You uploaded **.{ext}**.</div>",
                unsafe_allow_html=True,
            )

# =========================
# Branding & lists (expander)
# =========================
with st.expander("Branding & lists (optional)", expanded=False):
    # Logo uploader
    st.subheader("Logo", divider="gray")
    st.caption("Upload a **logo** (PNG/JPG/SVG). The banner updates immediately.")
    logo_up = st.file_uploader("Drag & drop logo here", key="logo_upl")
    st.session_state.logo_warning = ""
    if logo_up is not None:
        ext = Path(logo_up.name).suffix.lower().lstrip(".")
        allowed_ext = {"png", "jpg", "jpeg", "svg"}
        if ext in allowed_ext:
            st.session_state.logo_b64 = base64.b64encode(logo_up.getvalue()).decode("utf-8")
            st.session_state.logo_file_info = {"name": logo_up.name, "size": logo_up.size}
            st.session_state.logo_uploaded = True
            render_topbar(st.session_state.logo_b64)
        else:
            st.session_state.logo_uploaded = False
            st.session_state.logo_warning = f"Only PNG, JPG or SVG are supported for the banner logo. You uploaded **.{ext}**."
    if st.session_state.logo_warning:
        st.markdown(f"<div class='upload-warn'>⚠️ {st.session_state.logo_warning}</div>", unsafe_allow_html=True)
    if st.session_state.get("logo_uploaded"):
        info = st.session_state.get("logo_file_info", {})
        name = info.get("name", "logo")
        size_kb = round((info.get("size", 0) or 0) / 1024)
        st.markdown(
            f"""
            <span class="upload-chip">
              <svg viewBox="0 0 24 24" fill="none"><path d="M20 7L9 18l-5-5" stroke="white" stroke-width="2"
              stroke-linecap="round" stroke-linejoin="round"/></svg>
              Uploaded: {name} ({size_kb} KB)
            </span>
            """,
            unsafe_allow_html=True,
        )
        st.markdown("<div class='upload-note'>Your logo is now active in the top banner.</div>", unsafe_allow_html=True)

    # Courses list: show uploader only if NOT assets-managed
    st.subheader("Courses list", divider="gray")
    if not HAS_ASSET_COURSES:
        st.caption("Upload **courses.csv** with headers `code,label` to add/replace courses (no redeploy).")
        csv_up = st.file_uploader("Drag & drop courses.csv here", type=["csv"], key="courses_upl")
        c1, c2 = st.columns([1,2])
        with c1:
            st.download_button("Download CSV template", data=make_courses_template(),
                               file_name="courses_template.csv", type="secondary")
        with c2:
            st.caption("Tip: Put your full list in `/assets/courses.csv` to load automatically on start.")
        if csv_up is not None:
            try:
                reader = csv.DictReader(StringIO(csv_up.getvalue().decode("utf-8")))
                new_courses = []
                for r in reader:
                    code = (r.get("code") or "").strip()
                    label = (r.get("label") or "").strip()
                    if code and label:
                        new_courses.append((code, label))
                if new_courses:
                    set_courses(new_courses)
                    st.session_state.courses_uploaded = True
                    st.session_state.courses_file_info = {"name": csv_up.name, "count": len(new_courses)}
                else:
                    st.session_state.courses_uploaded = False
                    st.warning("No valid rows found. Expecting headers `code,label`.")
            except Exception as e:
                st.session_state.courses_uploaded = False
                st.error(f"Could not parse CSV: {e}")
        if st.session_state.get("courses_uploaded"):
            info = st.session_state.get("courses_file_info", {})
            name = info.get("name", "courses.csv"); count = info.get("count", 0)
            st.markdown(
                f"""
                <span class="upload-chip">
                  <svg viewBox="0 0 24 24" fill="none"><path d="M20 7L9 18l-5-5" stroke="white" stroke-width="2"
                  stroke-linecap="round" stroke-linejoin="round"/></svg>
                  Uploaded: {name} — {count} course(s)
                </span>
                """,
                unsafe_allow_html=True,
            )
            st.markdown("<div class='upload-note'>Course dropdown updated with your CSV.</div>", unsafe_allow_html=True)
    else:
        st.caption("Courses are managed in <code>/assets/courses.csv</code>. Edit there to update.",
                   unsafe_allow_html=True)

# =========================
# Static lists
# =========================
COHORTS = [
    "D1-C01","D1-E01","D1-E02","D1-M01","D1-M02","D1-M03","D1-M04","D1-M05",
    "D2-C01","D2-M01","D2-M02","D2-M03","D2-M04","D2-M05","D2-M06"
]
INSTRUCTORS = [
    "Ben","Abdulmalik","Gerhard","Faiz Lazam","Mohammed Alfarhan","Nerdeen","Dari","Ghamza",
    "Michail","Meshari","Mohammed Alwuthaylah","Myra","Meshal","Ibrahim","Khalil","Salem",
    "Rana","Daniel","Ahmed Albader"
]
VERBS = {
    "Low":    ["define","identify","list","recall","describe","classify","match"],
    "Medium": ["apply","solve","calculate","compare","analyze","demonstrate","explain"],
    "High":   ["evaluate","synthesize","design","justify","critique","optimize","create"]
}
def bloom_from_week(week: int) -> str:
    return "Low" if week <= 4 else ("Medium" if week <= 9 else "High")

# =========================
# Setup Row
# =========================
codes = course_codes()
labels = code_to_label()
if not st.session_state.course_code and codes:
    st.session_state.course_code = codes[0]

# Yellow tip only when truly on fallback and not assets/CSV uploaded
if len(codes) <= 6 and not (HAS_ASSET_COURSES or st.session_state.get("courses_uploaded")):
    st.markdown(
        "<div class='soft-tip'>Showing default courses. "
        "Upload a <strong>courses.csv</strong> in the Branding &amp; lists panel to add the rest.</div>",
        unsafe_allow_html=True
    )

st.markdown("<div class='adi-card'>", unsafe_allow_html=True)
r1c = st.columns([2, 1.8, .8, .8, 1.6])

with r1c[0]:
    display = [f"{c} — {labels.get(c,'')}" for c in codes]
    try:
        idx = codes.index(st.session_state.course_code)
    except ValueError:
        idx = 0
    sel = st.selectbox("Course (code)", options=list(range(len(codes))),
                       format_func=lambda i: display[i], index=idx, key="course_idx")
    st.session_state.course_code = codes[sel]
    long_name = labels.get(st.session_state.course_code, "")
    st.markdown(
        f"<span class='muted'>{long_name}</span><span class='pill pill-green'>{st.session_state.course_code}</span>",
        unsafe_allow_html=True
    )

with r1c[1]:
    st.selectbox("Class / Cohort", COHORTS,
                 index=COHORTS.index(st.session_state.class_cohort) if st.session_state.class_cohort in COHORTS else 0,
                 key="class_cohort")
    st.markdown(f"<span class='pill pill-green'>{st.session_state.class_cohort}</span>", unsafe_allow_html=True)

with r1c[2]:
    st.number_input("Lesson", min_value=1, max_value=20, step=1, key="lesson")
    st.markdown(f"<span class='pill pill-green'>L{int(st.session_state.lesson)}</span>", unsafe_allow_html=True)

with r1c[3]:
    def on_week_change():
        new_bloom = bloom_from_week(int(st.session_state.week))
        st.session_state.bloom_level = new_bloom
        st.session_state.verbs_selected = VERBS[new_bloom][:]
    st.number_input("Week", min_value=1, max_value=14, step=1, key="week", on_change=on_week_change)
    st.markdown(f"<span class='pill pill-green'>W{int(st.session_state.week)}</span>", unsafe_allow_html=True)

with r1c[4]:
    st.selectbox("Instructor", INSTRUCTORS, key="instructor")
    st.markdown(f"<span class='pill pill-green'>{st.session_state.instructor}</span>", unsafe_allow_html=True)

st.markdown("</div>", unsafe_allow_html=True)

# =========================
# Authoring
# =========================
st.markdown("### Authoring")
st.markdown("<div class='adi-card tight'>", unsafe_allow_html=True)

st.text_input("Topic / Outcome (optional)", key="topic_outcome", placeholder="e.g., Integrated Project and …")
st.caption(
    "ADI policy: Weeks 1–4 Low • 5–9 Medium • 10–14 High  |  "
    f"Recommended Bloom: **{bloom_from_week(int(st.session_state.week))}**"
)

st.segmented_control("Mode", ["Knowledge","Skills","Revision","Print Summary"], key="mode")

def update_verbs_on_bloom_change():
    allowed = set(VERBS[st.session_state.bloom_level])
    st.session_state.verbs_selected = [v for v in st.session_state.verbs_selected if v in allowed]

if st.session_state.mode != "Print Summary":
    a1, a2, a3, a4 = st.columns([1.2, .9, .9, 1.8])
    with a1:
        st.segmented_control("Bloom level", ["Low", "Medium", "High"],
                             key="bloom_level", on_change=update_verbs_on_bloom_change)
    with a2:
        st.button("Select all", on_click=lambda: st.session_state.update(
            {"verbs_selected": VERBS[st.session_state.bloom_level][:]}))
    with a3:
        st.button("Clear", on_click=lambda: st.session_state.update({"verbs_selected": []}))
    with a4:
        st.button("Use recommended for this week", on_click=lambda: st.session_state.update(
            {"bloom_level": bloom_from_week(int(st.session_state.week)),
             "verbs_selected": VERBS[bloom_from_week(int(st.session_state.week))][:]}
        ))

    verbs_for_level = VERBS[st.session_state.bloom_level]
    st.multiselect(f"Learning verbs (selected {len(st.session_state.verbs_selected)})",
                   options=verbs_for_level, key="verbs_selected")

    st.text_area("Topics (one per line)", key="topics_text", height=110,
                 placeholder="e.g.\n- Welding safety checks\n- NDT techniques (PT, MT, UT)\n- Inspection documentation flow")

    c1, c2, c3 = st.columns([1,1,2])
    with c1:
        include_key = st.checkbox("Answer key", value=True)
    with c2:
        mcq_count = st.selectbox("MCQs", [5,10,15,20], index=1)
    with c3:
        st.markdown("&nbsp;", unsafe_allow_html=True)
        if st.button("Generate MCQs", type="primary"):
            topics = [t.strip() for t in st.session_state.topics_text.splitlines() if t.strip()]
            topic0 = topics[0] if topics else "topic"
            st.session_state.generated_items = [{
                "stem": f"Sample question {i+1} on {topic0}?",
                "options": ["A) …", "B) …", "C) …", "D) …"],
                "answer": "A"
            } for i in range(int(mcq_count))]

    if st.session_state.generated_items:
        st.markdown("#### Preview")
        for idx, q in enumerate(st.session_state.generated_items):
            with st.expander(f"Q{idx+1}: {q['stem'][:90]}"):
                q["stem"] = st.text_input("Stem", value=q["stem"], key=f"stem-{idx}")
                a,b = st.columns(2)
                q["options"][0] = a.text_input("Option A", value=q["options"][0], key=f"oa-{idx}")
                q["options"][1] = b.text_input("Option B", value=q["options"][1], key=f"ob-{idx}")
                c,d = st.columns(2)
                q["options"][2] = c.text_input("Option C", value=q["options"][2], key=f"oc-{idx}")
                q["options"][3] = d.text_input("Option D", value=q["options"][3], key=f"od-{idx}")
                q["answer"] = st.selectbox("Correct", ["A","B","C","D"],
                                           index=["A","B","C","D"].index(q["answer"]), key=f"ans-{idx}")

        # TXT export
        txt = "\n\n".join(
            [f"Q{n+1}. {q['stem']}\n" + "\n".join(q["options"]) +
             (f"\nAnswer: {q['answer']}" if include_key else "")
             for n,q in enumerate(st.session_state.generated_items)]
        )
        st.download_button(
            "Export (TXT)", data=txt,
            file_name=f"{st.session_state.course_code}_L{st.session_state.lesson}_W{st.session_state.week}_mcqs.txt"
        )

        # DOCX export
        def as_docx(items) -> bytes | None:
            try:
                from docx import Document
                from docx.shared import Pt
                doc = Document()
                doc.add_heading(
                    f"{st.session_state.course_code} — Lesson {st.session_state.lesson} (Week {st.session_state.week})",
                    level=1
                )
                doc.add_paragraph()
                for i, q in enumerate(items, start=1):
                    doc.add_paragraph(f"Q{i}. {q['stem']}")
                    for opt in q["options"]:
                        doc.add_paragraph(opt)
                    if include_key:
                        p = doc.add_paragraph(f"Answer: {q['answer']}")
                        p.runs[0].font.bold = True
                    doc.add_paragraph()
                doc.styles["Normal"].font.name = "Calibri"
                doc.styles["Normal"].font.size = Pt(11)
                buf = BytesIO(); doc.save(buf); return buf.getvalue()
            except Exception:
                return None

        docx_bytes = as_docx(st.session_state.generated_items)
        if docx_bytes:
            st.download_button(
                "Export (Word .docx)", data=docx_bytes,
                file_name=f"{st.session_state.course_code}_L{st.session_state.lesson}_W{st.session_state.week}_mcqs.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.caption("Install `python-docx` for Word export (TXT is always available).")

else:
    # Print Summary
    ss = st.session_state
    topics_list = [t for t in ss.topics_text.splitlines() if t.strip()] or ["(add topics in other modes)"]
    st.markdown(
        f"""
        <div class="adi-card" style="max-width:860px;">
          <h2 style="margin:0 0 .3rem 0; color:{ADI_GREEN}">{ss.course_code} — Lesson {ss.lesson} (Week {ss.week})</h2>
          <div class="muted">{code_to_label().get(ss.course_code,'')}</div>
          <div class="muted"><strong>Instructor:</strong> {ss.instructor} &nbsp;|&nbsp;
               <strong>Class:</strong> {ss.class_cohort} &nbsp;|&nbsp;
               <strong>Bloom:</strong> {bloom_from_week(int(ss.week))}</div>
          <h3 style="margin:.8rem 0 .4rem 0;">Topics</h3>
          <ol style="margin-top:0;">{"".join(f"<li>{t}</li>" for t in topics_list)}</ol>
          <h3 style="margin:.8rem 0 .4rem 0;">MCQs (summary)</h3>
          <ol>{"".join(f"<li>{q['stem']}</li>" for q in (ss.generated_items or []))}</ol>
        </div>
        """,
        unsafe_allow_html=True,
    )

st.markdown("</div>", unsafe_allow_html=True)

