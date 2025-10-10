# export.py — TXT and DOCX export functions
from io import BytesIO
from docx import Document

def mcqs_to_txt(mcqs):
    lines = []
    for i, q in enumerate(mcqs, start=1):
        lines.append(f"Q{i}. {q['stem']}")
        lines.append(f"A. {q['A']}")
        lines.append(f"B. {q['B']}")
        lines.append(f"C. {q['C']}")
        lines.append(f"D. {q['D']}")
        lines.append(f"Answer: {q['correct']}")
        lines.append("")
    return "\n".join(lines)

def mcqs_to_docx_bytes(mcqs):
    doc = Document()
    doc.add_heading("ADI MCQs", level=1)
    for i, q in enumerate(mcqs, start=1):
        doc.add_paragraph(f"Q{i}. {q['stem']}")
        doc.add_paragraph(f"A. {q['A']}")
        doc.add_paragraph(f"B. {q['B']}")
        doc.add_paragraph(f"C. {q['C']}")
        doc.add_paragraph(f"D. {q['D']}")
        doc.add_paragraph(f"Answer: {q['correct']}")
        doc.add_paragraph("")
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()
